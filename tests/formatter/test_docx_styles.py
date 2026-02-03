from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from formatter.config import BodyStyle, FormatConfig, HeadingStyle
from formatter.docx_builder import build_docx


def test_docx_styles_and_margins_applied(tmp_path):
    config = FormatConfig()
    output = tmp_path / "out.docx"
    ast = [
        {"type": "heading", "level": 1, "text": "Title"},
        {"type": "paragraph", "text": "Hello"},
    ]
    build_docx(ast, output, config)

    doc = Document(output)
    section = doc.sections[0]
    assert round(section.top_margin.cm, 2) == 2.54
    assert round(section.left_margin.cm, 2) == 3.18

    normal = doc.styles["Normal"]
    assert normal.font.size == Pt(config.body_style.size_pt)

    h1 = doc.styles["Heading 1"]
    assert h1.font.size == Pt(config.heading_styles[1].size_pt)


def test_body_paragraph_indent_and_alignment(tmp_path):
    config = FormatConfig()
    output = tmp_path / "out.docx"
    build_docx([{"type": "paragraph", "text": "Hello"}], output, config)

    doc = Document(output)
    paragraph = doc.paragraphs[0]
    assert paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    assert paragraph.paragraph_format.first_line_indent == Pt(
        config.body_style.size_pt * 2
    )


def test_build_docx_falls_back_on_style_failure(tmp_path, monkeypatch):
    def _boom(*args, **kwargs):
        raise RuntimeError("style write failed")

    monkeypatch.setattr("formatter.docx_builder._apply_style_paragraph", _boom)

    output = tmp_path / "out.docx"
    config = FormatConfig()
    build_docx([{"type": "paragraph", "text": "Hello"}], output, config)

    doc = Document(output)
    paragraph = doc.paragraphs[0]
    assert paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    assert paragraph.paragraph_format.first_line_indent == Pt(
        config.body_style.size_pt * 2
    )


def test_docx_applies_line_spacing_and_indents(tmp_path):
    config = FormatConfig(
        body_style=BodyStyle(
            cn_font="SimSun",
            en_font="Times New Roman",
            size_pt=12,
            line_spacing=1.25,
            para_before_lines=0.0,
            para_after_lines=0.0,
            indent_before_chars=1,
            indent_after_chars=2,
            first_line_indent_chars=2,
            justify=True,
        ),
        heading_styles={
            1: HeadingStyle(
                font="SimHei",
                size_pt=14,
                line_spacing=1.25,
                para_before_lines=0.5,
                para_after_lines=0.5,
            )
        },
    )
    output = tmp_path / "out.docx"
    build_docx(
        [
            {"type": "heading", "level": 1, "text": "Title"},
            {"type": "paragraph", "text": "Hello"},
        ],
        output,
        config,
    )

    doc = Document(output)
    normal = doc.styles["Normal"]
    assert normal.paragraph_format.left_indent == Pt(12)
    assert normal.paragraph_format.right_indent == Pt(24)
    assert normal.paragraph_format.first_line_indent == Pt(24)

    h1 = doc.styles["Heading 1"]
    assert h1.paragraph_format.space_before == Pt(7)
    assert h1.paragraph_format.space_after == Pt(7)


def test_docx_heading_color_is_black(tmp_path):
    config = FormatConfig(
        heading_styles={
            1: HeadingStyle(
                font="SimHei",
                size_pt=14,
                line_spacing=1.25,
                para_before_lines=0.5,
                para_after_lines=0.5,
            )
        }
    )
    output = tmp_path / "out.docx"
    build_docx([{"type": "heading", "level": 1, "text": "标题"}], output, config)

    doc = Document(output)
    h1 = doc.styles["Heading 1"]
    assert h1.font.color.rgb == RGBColor(0, 0, 0)


def test_list_indents_and_numbering_levels(tmp_path):
    ast = [
        {
            "type": "list",
            "ordered": True,
            "level": 1,
            "start": 1,
            "items": [
                [{"type": "paragraph", "text": "第一项"}],
                [
                    {
                        "type": "list",
                        "ordered": True,
                        "level": 2,
                        "start": 1,
                        "items": [[{"type": "paragraph", "text": "子项"}]],
                    }
                ],
            ],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    first = doc.paragraphs[0]
    nested = doc.paragraphs[1]
    assert first.paragraph_format.left_indent == Pt(18)
    assert nested.paragraph_format.left_indent == Pt(36)
