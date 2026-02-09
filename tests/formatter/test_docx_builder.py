from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH

from formatter.config import FormatConfig
from formatter.docx_builder import build_docx


def test_build_docx_creates_docx(tmp_path):
    ast = [
        {"type": "heading", "level": 1, "text": "Title"},
        {"type": "paragraph", "text": "Hello"},
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())
    assert output.exists()


def test_build_docx_renders_bold_runs(tmp_path):
    ast = [
        {
            "type": "paragraph",
            "text": "Hello Bold",
            "runs": [
                {"text": "Hello ", "bold": False},
                {"text": "Bold", "bold": True},
            ],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())
    assert output.exists()

    doc = Document(output)
    runs = doc.paragraphs[0].runs
    assert runs[0].text == "Hello "
    assert runs[1].text == "Bold"
    assert runs[1].bold is True


def test_build_docx_renders_inline_styles(tmp_path):
    ast = [
        {
            "type": "paragraph",
            "runs": [
                {
                    "text": "It",
                    "bold": False,
                    "italic": True,
                    "strike": False,
                    "highlight": False,
                    "superscript": False,
                    "subscript": False,
                    "code": False,
                },
                {
                    "text": "S",
                    "bold": False,
                    "italic": False,
                    "strike": True,
                    "highlight": False,
                    "superscript": False,
                    "subscript": False,
                    "code": False,
                },
                {
                    "text": "H",
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "highlight": True,
                    "superscript": False,
                    "subscript": False,
                    "code": False,
                },
                {
                    "text": "2",
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "highlight": False,
                    "superscript": True,
                    "subscript": False,
                    "code": False,
                },
                {
                    "text": "2",
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "highlight": False,
                    "superscript": False,
                    "subscript": True,
                    "code": False,
                },
                {
                    "text": "code",
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "highlight": False,
                    "superscript": False,
                    "subscript": False,
                    "code": True,
                },
            ],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    runs = doc.paragraphs[0].runs
    assert runs[0].italic is True
    assert runs[1].font.strike is True
    assert runs[2].font.highlight_color == WD_COLOR_INDEX.YELLOW
    assert runs[3].font.superscript is True
    assert runs[4].font.subscript is True


def test_build_docx_renders_inline_code(tmp_path):
    ast = [
        {
            "type": "paragraph",
            "runs": [
                {
                    "text": "code",
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "highlight": False,
                    "superscript": False,
                    "subscript": False,
                    "code": True,
                }
            ],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    run = doc.paragraphs[0].runs[0]
    assert run.font.name == "Consolas"
    assert run.font.highlight_color == WD_COLOR_INDEX.GRAY_25


def test_build_docx_renders_lists(tmp_path):
    ast = [
        {
            "type": "list",
            "ordered": False,
            "level": 1,
            "start": 1,
            "items": [
                [
                    {
                        "type": "paragraph",
                        "text": "第一项",
                        "runs": [
                            {
                                "text": "第一项",
                                "bold": False,
                                "italic": False,
                                "strike": False,
                                "highlight": False,
                                "superscript": False,
                                "subscript": False,
                                "code": False,
                            }
                        ],
                    }
                ],
                [
                    {
                        "type": "paragraph",
                        "text": "第二项",
                        "runs": [
                            {
                                "text": "第二项",
                                "bold": False,
                                "italic": False,
                                "strike": False,
                                "highlight": False,
                                "superscript": False,
                                "subscript": False,
                                "code": False,
                            }
                        ],
                    }
                ],
            ],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())
    doc = Document(output)
    style = doc.paragraphs[0].style
    assert style is not None
    assert style.name and style.name.startswith("List Bullet")


def test_build_docx_renders_tables(tmp_path):
    ast = [
        {
            "type": "table",
            "align": ["left", "right"],
            "header": [
                {
                    "text": "列1",
                    "runs": [
                        {
                            "text": "列1",
                            "bold": False,
                            "italic": False,
                            "strike": False,
                            "highlight": False,
                            "superscript": False,
                            "subscript": False,
                            "code": False,
                        }
                    ],
                },
                {
                    "text": "列2",
                    "runs": [
                        {
                            "text": "列2",
                            "bold": False,
                            "italic": False,
                            "strike": False,
                            "highlight": False,
                            "superscript": False,
                            "subscript": False,
                            "code": False,
                        }
                    ],
                },
            ],
            "rows": [
                [
                    {
                        "text": "A",
                        "runs": [
                            {
                                "text": "A",
                                "bold": False,
                                "italic": False,
                                "strike": False,
                                "highlight": False,
                                "superscript": False,
                                "subscript": False,
                                "code": False,
                            }
                        ],
                    },
                    {
                        "text": "B",
                        "runs": [
                            {
                                "text": "B",
                                "bold": False,
                                "italic": False,
                                "strike": False,
                                "highlight": False,
                                "superscript": False,
                                "subscript": False,
                                "code": False,
                            }
                        ],
                    },
                ]
            ],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())
    doc = Document(output)
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "列1"
    assert doc.tables[0].cell(1, 1).text == "B"


def test_build_docx_renders_math(tmp_path):
    ast = [
        {"type": "paragraph", "runs": [{"type": "math", "latex": "a^2 + b^2 = c^2"}]},
        {"type": "math_block", "latex": "x"},
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())
    doc = Document(output)
    xml = doc.part._element.xml
    assert "oMath" in xml


def test_math_block_adds_equation_number(tmp_path):
    ast = [{"type": "math_block", "latex": "x"}]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    paragraph = doc.paragraphs[0]
    xml = paragraph._p.xml
    assert "SEQ Equation" in xml
    assert "oMath" in xml


def test_math_block_uses_complex_field_code(tmp_path):
    ast = [{"type": "math_block", "latex": "x"}]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    xml = doc.paragraphs[0]._p.xml
    assert 'w:fldChar w:fldCharType="begin"' in xml
    assert "w:instrText" in xml
    assert "SEQ Equation" in xml
    assert 'w:fldChar w:fldCharType="end"' in xml
    assert "w:fldSimple" not in xml


def test_math_blocks_number_increment(tmp_path):
    ast = [
        {"type": "math_block", "latex": "x"},
        {"type": "math_block", "latex": "y"},
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    xml = doc.part._element.xml
    assert xml.count("SEQ Equation") == 2


def test_math_block_in_list_keeps_numbering(tmp_path):
    ast = [
        {
            "type": "list",
            "ordered": False,
            "level": 1,
            "start": 1,
            "items": [[{"type": "math_block", "latex": "a"}]],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    para = doc.paragraphs[0]
    assert "SEQ Equation" in para._p.xml
    assert "oMath" in para._p.xml


def test_math_block_uses_center_and_right_tabs(tmp_path):
    ast = [{"type": "math_block", "latex": "x"}]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    xml = doc.paragraphs[0]._p.xml
    assert 'w:tab w:val="center"' in xml
    assert 'w:tab w:val="right"' in xml
    assert "<w:tab/>" in xml
    assert xml.find("<w:tab/>") < xml.find("oMath")


def test_table_cells_are_center_aligned(tmp_path):
    ast = {
        "type": "table",
        "align": ["left", "right"],
        "header": [{"text": "H1"}, {"text": "H2"}],
        "rows": [[{"text": "A"}, {"text": "B"}]],
    }
    output = tmp_path / "out.docx"
    build_docx([ast], output, FormatConfig())

    doc = Document(output)
    table = doc.tables[0]
    for row in table.rows:
        for cell in row.cells:
            assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_table_cells_clear_first_line_indent(tmp_path):
    ast = {
        "type": "table",
        "align": ["left"],
        "header": [{"text": "H1"}],
        "rows": [[{"text": "A"}]],
    }
    output = tmp_path / "out.docx"
    build_docx([ast], output, FormatConfig())

    doc = Document(output)
    para = doc.tables[0].cell(0, 0).paragraphs[0]
    xml = para._p.xml
    assert 'w:firstLine="0"' in xml


def test_table_cell_text_trims_leading_spaces(tmp_path):
    ast = {
        "type": "table",
        "align": ["left"],
        "header": [{"text": "H1", "runs": [{"text": "  H1", "code": False}]}],
        "rows": [[{"text": "A", "runs": [{"text": "  A", "code": False}]}]],
    }
    output = tmp_path / "out.docx"
    build_docx([ast], output, FormatConfig())

    doc = Document(output)
    table = doc.tables[0]
    assert table.cell(0, 0).text == "H1"
    assert table.cell(1, 0).text == "A"


def test_table_uses_three_line_style(tmp_path):
    ast = [
        {
            "type": "table",
            "align": ["left"],
            "header": [{"text": "列1"}],
            "rows": [[{"text": "A"}]],
        }
    ]
    output = tmp_path / "out.docx"
    build_docx(ast, output, FormatConfig())

    doc = Document(output)
    tbl_xml = doc.tables[0]._tbl.xml
    assert "w:tblBorders" in tbl_xml
    assert 'w:top w:val="single" w:sz="12"' in tbl_xml
    assert 'w:bottom w:val="single" w:sz="12"' in tbl_xml
    assert 'w:bottom w:val="single" w:sz="6"' in tbl_xml
    assert "insideV" not in tbl_xml
