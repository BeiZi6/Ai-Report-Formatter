from __future__ import annotations

from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from formatter.config import FormatConfig
from formatter.latex import latex_to_omml


def _set_style_fonts(style, ascii_font: str, east_asia_font: str | None = None) -> None:
    style.font.name = ascii_font
    rfonts = style.element.rPr.rFonts
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    if east_asia_font:
        rfonts.set(qn("w:eastAsia"), east_asia_font)


def _lines_to_pt(lines: float, font_size_pt: int) -> Pt:
    return Pt(lines * font_size_pt)


def _chars_to_pt(chars: int, font_size_pt: int) -> Pt:
    return Pt(chars * font_size_pt)


LIST_INDENT_PT = 18
THREE_LINE_BORDER_THICK_SZ = 12
THREE_LINE_BORDER_THIN_SZ = 6


def _apply_list_indents(paragraph, level: int) -> None:
    paragraph.paragraph_format.left_indent = Pt(LIST_INDENT_PT * level)
    paragraph.paragraph_format.first_line_indent = Pt(-LIST_INDENT_PT / 2)


def _apply_style_paragraph(style, size_pt, line_spacing, before_lines, after_lines, align=None) -> None:
    style.font.size = Pt(size_pt)
    pf = style.paragraph_format
    pf.space_before = _lines_to_pt(before_lines, size_pt)
    pf.space_after = _lines_to_pt(after_lines, size_pt)
    pf.line_spacing = line_spacing
    if align is not None:
        pf.alignment = align


def _apply_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:shd"))
    if existing is not None:
        p_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _add_page_number(section, position: str) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if position == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def _add_math_run(paragraph, latex: str) -> None:
    run = paragraph.add_run()
    try:
        omml = latex_to_omml(latex)
        omml = _ensure_omml_namespace(omml)
        run._r.append(parse_xml(omml))
    except Exception:
        run.text = latex


def _apply_equation_tabs(paragraph, center_pos: int, right_pos: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:tabs"))
    if existing is not None:
        p_pr.remove(existing)
    tabs = OxmlElement("w:tabs")
    center_tab = OxmlElement("w:tab")
    center_tab.set(qn("w:val"), "center")
    center_tab.set(qn("w:pos"), str(center_pos))
    tabs.append(center_tab)
    right_tab = OxmlElement("w:tab")
    right_tab.set(qn("w:val"), "right")
    right_tab.set(qn("w:pos"), str(right_pos))
    tabs.append(right_tab)
    p_pr.append(tabs)


def _append_field_char(paragraph, field_type: str, dirty: bool = False) -> None:
    run = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), field_type)
    if dirty:
        fld.set(qn("w:dirty"), "true")
    run.append(fld)
    paragraph._p.append(run)


def _append_instr_text(paragraph, instruction: str) -> None:
    run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.text = instruction
    run.append(instr)
    paragraph._p.append(run)


def _append_text_run(paragraph, text: str) -> None:
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    paragraph._p.append(run)


def _add_equation_number_field(paragraph) -> None:
    paragraph.add_run("(")
    _append_field_char(paragraph, "begin", dirty=True)
    _append_instr_text(paragraph, r"SEQ Equation \* ARABIC")
    _append_field_char(paragraph, "separate")
    _append_text_run(paragraph, "1")
    _append_field_char(paragraph, "end")
    paragraph.add_run(")")


def _apply_run_styles(docx_run, run: dict) -> None:
    docx_run.bold = bool(run.get("bold"))
    docx_run.italic = bool(run.get("italic"))
    if run.get("strike"):
        docx_run.font.strike = True
    if run.get("highlight"):
        docx_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if run.get("superscript"):
        docx_run.font.superscript = True
    if run.get("subscript"):
        docx_run.font.subscript = True
    if run.get("code"):
        docx_run.font.name = "Consolas"
        if not run.get("highlight"):
            docx_run.font.highlight_color = WD_COLOR_INDEX.GRAY_25


def _ensure_omml_namespace(omml: str) -> str:
    if "xmlns:m=" in omml:
        return omml
    if "<m:oMathPara" in omml:
        return omml.replace(
            "<m:oMathPara",
            '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
            1,
        )
    return omml.replace(
        "<m:oMath",
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
        1,
    )


def _add_runs(paragraph, runs: list[dict], fallback_text: str = "") -> None:
    if not runs:
        if fallback_text:
            paragraph.add_run(fallback_text)
        return
    for run in runs:
        if run.get("type") == "math":
            _add_math_run(paragraph, run.get("latex", ""))
            continue
        text = run.get("text", "")
        if not text:
            continue
        docx_run = paragraph.add_run(text)
        _apply_run_styles(docx_run, run)


def _trim_leading_text_runs(runs: list[dict]) -> list[dict]:
    if not runs:
        return runs
    trimmed_runs: list[dict] = []
    trimmed = False
    for run in runs:
        if run.get("type") == "math":
            trimmed_runs.append(run)
            continue
        text = run.get("text", "")
        if not trimmed:
            stripped = text.lstrip()
            if stripped:
                if stripped != text:
                    run = {**run, "text": stripped}
                trimmed = True
            else:
                if not text:
                    continue
                continue
        trimmed_runs.append(run)
    return trimmed_runs


def _add_code_block(doc, node: dict) -> None:
    paragraph = doc.add_paragraph("")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _apply_paragraph_shading(paragraph, "F2F2F2")

    code_font = "Consolas"
    lines = (node.get("text") or "").splitlines()
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = code_font
        run.font.size = Pt(10)


def _list_style_name(ordered: bool, level: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    if level <= 1:
        return base
    level_suffix = min(level, 3)
    return f"{base} {level_suffix}"


def _add_math_block(paragraph, latex: str, center_tab: int, right_tab: int) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _apply_equation_tabs(paragraph, center_tab, right_tab)
    paragraph.add_run("\t")
    _add_math_run(paragraph, latex)
    paragraph.add_run("\t")
    _add_equation_number_field(paragraph)


def _add_list(doc, node: dict, config: FormatConfig, center_tab: int, right_tab: int) -> None:
    for item in node.get("items", []):
        for child in item:
            if child.get("type") == "paragraph":
                style_name = _list_style_name(node.get("ordered", False), node.get("level", 1))
                paragraph = doc.add_paragraph("", style=style_name)
                _add_runs(paragraph, child.get("runs", []), child.get("text", ""))
                _apply_list_indents(paragraph, node.get("level", 1))
            elif child.get("type") == "list":
                _add_list(doc, child, config, center_tab, right_tab)
            elif child.get("type") == "math_block":
                paragraph = doc.add_paragraph("", style=_list_style_name(node.get("ordered", False), node.get("level", 1)))
                _add_math_block(paragraph, child.get("latex", ""), center_tab, right_tab)
                _apply_list_indents(paragraph, node.get("level", 1))
            elif child.get("type") == "table":
                _add_table(doc, child)
            elif child.get("type") == "code_block":
                _add_code_block(doc, child)


def _cell_alignment(align: str):
    if align == "center":
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    if align == "right":
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    return WD_PARAGRAPH_ALIGNMENT.LEFT


def _apply_three_line_table(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(THREE_LINE_BORDER_THICK_SZ))
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    for edge in ("left", "right", "insideH"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)

    inside_v = borders.find(qn("w:insideV"))
    if inside_v is not None:
        borders.remove(inside_v)


def _apply_header_bottom_border(row) -> None:
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(THREE_LINE_BORDER_THIN_SZ))
        bottom.set(qn("w:color"), "000000")
        tc_borders.append(bottom)


def _add_table(doc, node: dict) -> None:
    header = node.get("header", [])
    rows = node.get("rows", [])
    if not header:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    # Avoid built-in styles that reintroduce vertical borders
    try:
        table.style = None
    except Exception:
        table.style = "Table Grid"
    _apply_three_line_table(table)

    all_rows = [header] + rows
    for r_idx, row in enumerate(all_rows):
        for c_idx, cell in enumerate(row):
            cell_obj = table.cell(r_idx, c_idx)
            cell_obj.text = ""
            paragraph = cell_obj.paragraphs[0]
            runs = _trim_leading_text_runs(cell.get("runs", []))
            fallback_text = (cell.get("text") or "").lstrip()
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _add_runs(paragraph, runs, fallback_text)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if r_idx == 0:
            _apply_header_bottom_border(table.rows[r_idx])


def build_docx(ast: list[dict], output_path, config: FormatConfig | None = None) -> None:
    config = config or FormatConfig()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    try:
        _add_page_number(section, config.page_num_position)
    except Exception:
        pass

    try:
        normal = doc.styles["Normal"]
        _set_style_fonts(normal, config.body_style.en_font, config.body_style.cn_font)
        _apply_style_paragraph(
            normal,
            config.body_style.size_pt,
            config.body_style.line_spacing,
            config.body_style.para_before_lines,
            config.body_style.para_after_lines,
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY if config.body_style.justify else None,
        )
        paragraph_format = getattr(normal, "paragraph_format", None)
        if paragraph_format is not None:
            paragraph_format.left_indent = _chars_to_pt(
                config.body_style.indent_before_chars, config.body_style.size_pt
            )
            paragraph_format.right_indent = _chars_to_pt(
                config.body_style.indent_after_chars, config.body_style.size_pt
            )
            paragraph_format.first_line_indent = _chars_to_pt(
                config.body_style.first_line_indent_chars, config.body_style.size_pt
            )

        for level, hstyle in config.heading_styles.items():
            h = doc.styles[f"Heading {level}"]
            _set_style_fonts(h, hstyle.en_font, hstyle.cn_font)
            font = getattr(h, "font", None)
            if font is not None:
                font.color.rgb = RGBColor(0, 0, 0)
            _apply_style_paragraph(
                h,
                hstyle.size_pt,
                hstyle.line_spacing,
                hstyle.para_before_lines,
                hstyle.para_after_lines,
            )
    except Exception:
        pass

    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    if page_width and left_margin and right_margin:
        usable_width = page_width - left_margin - right_margin
        right_tab = int(getattr(usable_width, "twips", 9350))
        center_tab = int(right_tab / 2)
    else:
        right_tab = 9350
        center_tab = 4675

    for node in ast:
        if node.get("type") == "heading":
            paragraph = doc.add_heading("", level=node.get("level", 1))
            _add_runs(paragraph, node.get("runs", []), node.get("text", ""))
        elif node.get("type") == "paragraph":
            paragraph = doc.add_paragraph("")
            _add_runs(paragraph, node.get("runs", []), node.get("text", ""))
            if config.body_style.justify:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.left_indent = _chars_to_pt(
                config.body_style.indent_before_chars, config.body_style.size_pt
            )
            paragraph.paragraph_format.right_indent = _chars_to_pt(
                config.body_style.indent_after_chars, config.body_style.size_pt
            )
            paragraph.paragraph_format.first_line_indent = _chars_to_pt(
                config.body_style.first_line_indent_chars, config.body_style.size_pt
            )
        elif node.get("type") == "list":
            _add_list(doc, node, config, center_tab, right_tab)
        elif node.get("type") == "table":
            _add_table(doc, node)
        elif node.get("type") == "math_block":
            paragraph = doc.add_paragraph("")
            _add_math_block(paragraph, node.get("latex", ""), center_tab, right_tab)
        elif node.get("type") == "code_block":
            _add_code_block(doc, node)

    doc.save(output_path)
